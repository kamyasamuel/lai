# server.py
import logging
import sys
import json
import argparse
from openai import OpenAI
from groq import Groq
import os
import tornado.ioloop
from  tornado.web import Application, HTTPError # type: ignore
import tornado
import tornado.escape
from dotenv import load_dotenv
from googleapiclient.discovery import build
import tornado.websocket
from docx import Document
import PyPDF2
import glob  # Import the glob module
from urllib.parse import urlparse, parse_qs

# Load environment variables from .env file
load_dotenv()

groq_client = Groq(api_key=os.getenv('GROQ_API_KEY'))
openai_client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
deepseek_client = OpenAI(api_key=os.getenv("DEEPSEEK_API_KEY"), base_url="https://api.deepseek.com")

UPLOAD_DIR = './uploads'

# Initialize Langchain components for RAG
#embeddings = OpenAIEmbeddings()
# Placeholder for FAISS database - will be loaded/created later
# In a real application, you would build this index from your legal documents
# For now, we'll use a dummy Chroma DB for demonstration
#db = Chroma.from_texts(
#    ["This is a sample contract clause about termination.", "This is a sample regulation regarding data privacy in Nigeria.", "This is a summary of a case law on intellectual property."],
#    embeddings
#)

async def generate_draft(prompt: str) -> str:
    """
    Calls the OpenAI Chat API with a system prompt that tells the model
    to act as a legal assistant, then returns the drafted text.
    """
    messages = [
        {
            "role": "system",
            "content": (
                "You are an expert legal assistant. "
                "When given a user request, you will draft a clear, "
                "concise legal document or agreement that fulfills the requirements. "
                "Use proper legal structure, headings, and plain English."
                "Avoid ** at all cost"
            )
        },
        {
            "role": "user",
            "content": prompt
        }
    ]

    # call the OpenAI chat completion endpoint
    response = openai_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages, # type: ignore
        temperature=0.2,      # lower = more precise/legal
        max_tokens=1500       # adjust to taste
    )

    # extract the assistant’s reply
    draft_text = response.choices[0].message.content.strip() # type: ignore
    return draft_text

def get_document_text(file_data) -> str:
    """
    Extracts text content from docx or pdf file data.
    Employ OCD processing for pdf files.
    """
    filename = file_data['filename']
    file_body = file_data['body']

    text = ""
    if filename.lower().endswith(".docx"):
        # Create a BytesIO object from the file body
        from io import BytesIO
        byte_stream = BytesIO(file_body)
        document = Document(byte_stream)
        for paragraph in document.paragraphs:
            text += paragraph.text + ' '
    elif filename.lower().endswith(".pdf"):
        # Create a BytesIO object from the file body
        from io import BytesIO
        byte_stream = BytesIO(file_body)
        try:
            # Use PdfReader from PyPDF2
            reader = PyPDF2.PdfReader(byte_stream)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text() + ' '
        except Exception as e:
            logging.error(f"Error reading PDF file {filename}: {e}")
            raise HTTPError(400, f"Could not read PDF file {filename}.")
    else:
        raise HTTPError(400, f"Unsupported file type: {filename}. Only .docx and .pdf are supported.")

    return text

class BaseCORSHandler(tornado.web.RequestHandler):
    def set_default_headers(self):
        # Allow React dev server on port 5173
        self.set_header("Access-Control-Allow-Origin", "https://lawyers.legalaiafrica.com")
        self.set_header("Access-Control-Allow-Credentials", "true")
        self.set_header("Access-Control-Allow-Methods", "GET,POST,OPTIONS")
        self.set_header(
            "Access-Control-Allow-Headers",
            "Content-Type, Access-Control-Allow-Origin, Authorization"
        )

    def options(self):
        # No body for preflight `OPTIONS` requests
        self.set_status(204)
        self.finish()

class DraftHandler(BaseCORSHandler):
    async def post(self):
        try:
            data = tornado.escape.json_decode(self.request.body)
        except json.JSONDecodeError:
            self.set_status(400)
            return self.write({"error": "Invalid JSON"})

        prompt = data.get("prompt", "").strip()
        if not prompt:
            self.set_status(400)
            return self.write({"error": 'Missing "prompt" in request body'})

        draft_text = await generate_draft(prompt)

        self.set_header("Content-Type", "application/json")
        self.write({"draft": draft_text})

class AnalysisHandler(BaseCORSHandler):
    async def post(self):
         # Check if the upload directory exists, create if not
        if not os.path.exists(UPLOAD_DIR):
            os.makedirs(UPLOAD_DIR)

        # Get the files from the request
        files = self.request.files.get("file", [])

        if not files:
            raise HTTPError(400, "No file uploaded")

        for file_info in files:
            # Extract the filename and file body
            filename = file_info['filename']
            file_body = file_info['body']

            # Create a full path for the new file
            file_path = os.path.join(UPLOAD_DIR, filename)

            # Write the file to the specified path for backup purposes
            with open(file_path, 'wb') as f:
                f.write(file_body)

            text = ""
            if filename.endswith("docx"):
                document = Document(file_path)
                for paragraph in document.paragraphs:
                    text+=paragraph.text +''
            elif filename.endswith("pdf"):
                ...

            # Generate summary
            analysis = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role":"system", "content":"""You are an anlysis agent. You analyze documents for major key points and take-aways.
                     Format your output as html elements for insertion into the body section of an html file. 
                     Use headings, tables, pargraphs, sections, etc where possible.
                     Strictly output html elements only without extraneous tags like ```html or ```.
                    """},
                    {"role":"user", "content":f"Analyse the following document: {text}"}
                ],
                temperature=0.5,
                top_p=1
            )
        self.write(json.dumps({'summary': analysis.choices[0].message.content, 'filename':filename})) # type: ignore

class ChatWebSocketHandler(tornado.websocket.WebSocketHandler):
    def check_origin(self, origin):
        # allow React dev server or any origin you trust:
        return True

    async def on_message(self, message):
        """
        Expects a JSON payload:
          { "message": "Hello, please draft ..." }
        Streams back JSON messages:
          { "content": "First bit..." }
          { "content": "Next chunk..." }
          { "event": "done" }
        """
        try:
            data = json.loads(message)
            prompt = data.get("user_message")
            if not prompt:
                await self.write_message(json.dumps({"error":"Missing prompt"}))
                return

            # Stream chunks from OpenAI
            request = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                stream=True,
                messages=[
                    {"role":"system", "content":"""You are a chat agent for Legal Ai Africa.
                     Format your output into paragraphs.
                    """},
                    {"role": "user", "content":prompt}
                ]
            )

            for chunk in request:
                print(chunk)
                message = chunk.choices[0].delta.content
                await self.write_message(json.dumps({ "content": message }))

            # indicate end of stream
            await self.write_message(json.dumps({ "event": "done" }))

        except Exception as e:
            print(e)
            await self.write_message(json.dumps({
                "error": f"Server error: {e}"
            }))
            
class QueryHandler(BaseCORSHandler):
    async def post(self):
        try:
            data = json.loads(self.request.body)
            logging.info(data)
        except json.JSONDecodeError:
            self.set_status(400)
            return self.write({"error": "Invalid JSON"})

        query_type = data.get("type")
        query_text = data.get("query")

        if not query_type or not query_text:
            self.set_status(400)
            return self.write({"error": 'Missing "type" or "query" in request body'})

        response_text = ""
        try:
            if query_type in ["Contract Search", "Laws & Regulations", "Case Law"]:
                # Ensure the FAISS database is loaded
                global db
                if db is None:
                    # Assuming a pre-built FAISS index named "faiss_index" exists
                    # In a real application, you would build this index from your legal documents
                    try:
                        db = FAISS.load_local("faiss_index", embeddings)
                    except Exception as e:
                        logging.error(f"Error loading FAISS index: {e}")
                        self.set_status(500)
                        return self.write({"error": "FAISS index not available."})
                        
                retriever = db.as_retriever()
                
                # Use RetrievalQA chain for these types
                llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.2)
                qa_chain = RetrievalQA.from_chain_type(llm, retriever=retriever)
                result = qa_chain.invoke({"query": query_text})
                response_text = result.get("result", "No relevant information found.")

            elif query_type == "Web & News":
                # Use DuckDuckGo Search for Web & News
                search = duckduckgo_search()
                # Perform a search
                search_results = search.run(query_text)
                # You might want to process these results further with an LLM
                # For this example, we'll just return the raw search results
                response_text = search_results
            else:
                self.set_status(400)
                return self.write({"error": f"Unsupported query type: {query_type}"})

            self.set_header("Content-Type", "application/json")
            self.write({"response": response_text})
        except Exception as e:
            raise HTTPError(500, f"An error occurred during query processing: {e}")
            
class DocumentComparisonHandler(BaseCORSHandler):
    async def post(self):
        try:
            files = self.request.files
            if 'document1' not in files or 'document2' not in files:
                self.set_status(400)
                return self.write({"error": "Two document files named 'document1' and 'document2' are required."})

            doc1_file_data = files['document1'][0]
            doc2_file_data = files['document2'][0]

            # Extract text content from files
            doc1_content = get_document_text(doc1_file_data)
            doc2_content = get_document_text(doc2_file_data)

            # Use OpenAI API to compare documents
            comparison_response = openai_client.chat.completions.create(
                model="gpt-4o", # Or a suitable model for comparison
                messages=[
                    {"role": "system", "content": "You are an expert document comparison AI. Analyze the two provided documents and highlight key differences and similarities in a clear and concise manner. Format the output as a readable text."},
                    {"role": "user", "content": f"Compare Document 1 and Document 2. Document 1:{doc1_content} Document 2:{doc2_content}"}
                ],
                temperature=0.2,
                max_tokens=2000 # Adjust as needed
            )

            comparison_result = comparison_response.choices[0].message.content.strip()

            self.set_header("Content-Type", "application/json")
            self.write({"comparisonResult": comparison_result})

        except HTTPError as e:
            # Re-raise HTTPError to be handled by Tornado
            raise e
        except Exception as e:
            logging.error(f"Error during document comparison: {e}")
            self.set_status(500)
            self.write({"error": "Failed to compare documents using AI.", "details": str(e)})

class DocumentLibraryHandler(BaseCORSHandler):
    async def get(self):
        try:
            # Get the search term from the query string
            query = self.request.uri
            parsed_url = urlparse(query)
            query_params = parse_qs(parsed_url.query)
            search_term = query_params.get('search', [''])[0]  # Default to empty string if no search term

            # List all files in the UPLOAD_DIR
            all_files = glob.glob(os.path.join(UPLOAD_DIR, '*.*'))  # Match any file extension

            # Filter files based on the search term (if provided)
            if search_term:
                search_term_lower = search_term.lower()
                filtered_files = [f for f in all_files if os.path.basename(f).lower().find(search_term_lower) != -1]
            else:
                filtered_files = all_files

            # Format the file list for the response
            document_list = [os.path.basename(f) for f in filtered_files]

            self.set_header("Content-Type", "application/json")
            self.write(json.dumps(document_list))

        except Exception as e:
            logging.error(f"Error in document library handler: {e}")
            self.set_status(500)
            self.write({"error": "Failed to retrieve document list.", "details": str(e)})

class Application(Application): # type: ignore
    def __init__(self):
        handlers = [
            (r"/draft", DraftHandler),
            (r"/analyze/(.*)", AnalysisHandler),
            (r"/ws/chat", ChatWebSocketHandler),
            (r"/query", QueryHandler), # New query endpoint
            (r"/compare-documents", DocumentComparisonHandler), # New comparison endpoint
            (r"/documents", DocumentLibraryHandler), # New document library endpoint
        ]
        settings = {
            "debug": True,   # reload on change, more verbose errors
            "autoreload": True,
        }
        super().__init__(handlers, **settings)


def run_server(port: int = 4040):
    app = Application()
    app.listen(port)
    print(f"[server] Listening on http://localhost:{port}")
    tornado.ioloop.IOLoop.current().start()

def run_exec_mode(prompt: str):
    result = {"draft": generate_draft(prompt)}
    sys.stdout.write(json.dumps(result, indent=2))

def main():
    parser = argparse.ArgumentParser(description="Legal-AI Tornado Server")
    parser.add_argument(
        "--mode", choices=["server", "exec"], default="server",
        help="server = start HTTP server; exec = generate a draft and exit"
    )
    parser.add_argument(
        "--prompt", type=str,
        help="Prompt text (only used in --mode=exec)"
    )
    parser.add_argument(
        "--port", "-p", type=int, default=4040,
        help="Port to listen on (server mode)"
    )
    args = parser.parse_args()

    if args.mode == "exec":
        if not args.prompt:
            print("Error: --prompt is required in exec mode", file=sys.stderr)
            sys.exit(1)
        run_exec_mode(args.prompt)
    else:
        run_server(port=args.port)

if __name__ == "__main__":
    main()
